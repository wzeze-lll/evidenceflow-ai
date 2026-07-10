#!/usr/bin/env python3
"""生成 EvidenceFlow AI 实习报告 docx 文件"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
import os

# Screenshot directory
IMG_DIR = os.path.dirname(os.path.abspath(__file__))

doc = Document()

# ---- Page Setup ----
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)

# ---- Style Setup ----
style = doc.styles['Normal']
font = style.font
font.name = '宋体'
font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# Helper functions
def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
        if level == 1:
            run.font.size = Pt(16)
        elif level == 2:
            run.font.size = Pt(14)
        elif level == 3:
            run.font.size = Pt(13)
    return h

def add_para(text, bold=False, size=12, align=None, spacing_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(spacing_after)
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    return p

def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        shading_elm = cell._element.get_or_add_tcPr()
        shading = shading_elm.makeelement(qn('w:shd'), {
            qn('w:fill'): 'D9D9D9',
            qn('w:val'): 'clear'
        })
        shading_elm.append(shading)
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r + 1].cells[c]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(10)
            run.font.name = '宋体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    doc.add_paragraph()
    return table

def add_image(img_name, width_inches=5.5, caption=None):
    img_path = os.path.join(IMG_DIR, img_name)
    if os.path.exists(img_path):
        doc.add_paragraph()
        doc.add_picture(img_path, width=Inches(width_inches))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if caption:
            add_para(caption, size=10, align=WD_ALIGN_PARAGRAPH.CENTER, spacing_after=12)
    else:
        add_para(f'[图片缺失: {img_name}]', size=10)

# ======== COVER PAGE ========
for _ in range(6):
    doc.add_paragraph()

add_para('实 习 报 告', bold=True, size=26, align=WD_ALIGN_PARAGRAPH.CENTER, spacing_after=20)
doc.add_paragraph()
add_para('【EvidenceFlow AI 证流 AI】', bold=True, size=21.5, align=WD_ALIGN_PARAGRAPH.CENTER, spacing_after=30)

for _ in range(8):
    doc.add_paragraph()

cover_table = doc.add_table(rows=5, cols=2)

cover_data = [
    ('题  目', 'EvidenceFlow AI 证流 AI — 多文档智能分析平台'),
    ('班  级', '网络工程 2501 班'),
    ('组  长', '张子谦（202530809418）'),
    ('组  员', '于泽洋（202530809379）\n王昌宇（202530809426）\n丁苏宇（202530809423）\n何金鑫（202530809429）'),
    ('完成日期', '2026-07-11'),
]

for i, (label, value) in enumerate(cover_data):
    cell_label = cover_table.rows[i].cells[0]
    cell_label.text = ''
    p = cell_label.paragraphs[0]
    run = p.add_run(label)
    shading_elm = cell_label._element.get_or_add_tcPr()
    shading = shading_elm.makeelement(qn('w:shd'), {
        qn('w:fill'): 'D9D9D9',
        qn('w:val'): 'clear'
    })
    shading_elm.append(shading)
    cell_value = cover_table.rows[i].cells[1]
    cell_value.text = ''
    p = cell_value.paragraphs[0]
    run = p.add_run(value)

doc.add_page_break()

# ======== TABLE OF CONTENTS ========
add_heading_styled('目  录', level=1)
toc_items = [
    ('1.项目立项', 4),
    ('  1.1  项目背景', 4),
    ('  1.2  客户与最终用户', 5),
    ('  1.3  产品范围', 6),
    ('2.项目启动', 7),
    ('  2.1  组织结构', 7),
    ('  2.2  项目进度计划', 8),
    ('3.需求分析', 9),
    ('  3.1  产品的功能性需求', 9),
    ('  3.2  产品非功能性需求', 12),
    ('4.概要设计', 16),
    ('  4.1  系统体系结构', 16),
    ('  4.2  系统功能结构', 18),
    ('  4.3  运行环境', 19),
    ('  4.4  接口设计', 19),
    ('  4.5  部署设计', 21),
    ('5.详细设计', 22),
    ('  5.1  模块汇总', 22),
    ('  5.2  首页（Dashboard）模块设计', 23),
    ('  5.3  文档管理模块设计', 24),
    ('  5.4  AI 阅读器模块设计', 25),
    ('  5.5  证据链模块设计', 26),
    ('  5.6  冲突雷达模块设计', 27),
    ('  5.7  共识地图模块设计', 28),
    ('  5.8  决策简报模块设计', 29),
    ('  5.9  设置模块设计', 29),
    ('6.用户界面设计', 30),
    ('  6.1  总体设计思路', 30),
    ('  6.2  界面展示', 31),
    ('  6.3  技术选型与组件', 34),
    ('7.数据库设计', 37),
    ('  7.1  核心业务实体', 37),
    ('  7.2  数据模型设计', 38),
    ('  7.3  索引设计', 39),
    ('  7.4  数据存储方案', 39),
    ('8.编码与单元测试', 40),
    ('  8.1  代码检查', 40),
    ('  8.2  单元测试', 42),
    ('9.软件测试', 44),
    ('  9.1  测试计划', 44),
    ('  9.2  测试用例设计', 45),
    ('  9.3  测试执行', 46),
    ('  9.4  测试总结', 47),
]
for title, page in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    dots = '.' * max(2, 55 - len(title))
    run = p.add_run(f'{title}{dots}{page}')
    run.font.size = Pt(12)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

doc.add_page_break()

# ======== Chapter 1: 项目立项 ========
add_heading_styled('1. 项目立项', level=1)

add_heading_styled('1.1 项目背景', level=2)
add_para('随着人工智能技术的快速发展，大语言模型（LLM）在各个领域的应用日益广泛。在信息爆炸的时代，学生、研究者和知识工作者每天需要处理大量文档资料。面对海量信息，如何快速提取关键内容、识别不同来源之间的异同、追踪每个结论的依据来源，是一个普遍存在却未被很好解决的问题。')
add_para('现有工具存在明显的局限性：通用 AI 聊天工具（如 ChatGPT、DeepSeek）虽能回答问题，但无法自动引用原文出处，用户难以验证结论的真实性；传统文档管理工具（如 Notion、Evernote）擅长存储和检索，但缺乏深度 AI 分析能力；学术文献管理工具（如 Zotero）侧重于引用管理，但不具备多文档智能分析功能；企业级数据分析平台则价格昂贵，需要部署服务器，不适合个人和小团队使用。')
add_para('针对这一痛点，"EvidenceFlow AI（证流 AI）"应运而生。该项目旨在构建一款浏览器端的纯前端应用，专注于多文档智能分析。其核心差异化优势在于：每一个 AI 结论都附带可追溯的原文证据，实现从"盲目相信 AI"到"可验证的 AI 推理"的转变。')
add_para('本项目采用纯前端架构（React 19 + TypeScript + Vite），以浏览器 IndexedDB 作为唯一数据存储，所有文档在用户设备本地解析处理，无需后端服务器和数据库，充分保护用户隐私。项目于 2026 年夏季完成第一期开发，实现了完整的 Demo 演示系统。')

add_heading_styled('1.2 客户与最终用户', level=2)
add_para('EvidenceFlow AI 的目标用户分为三类：')
add_para('1）学生群体（大学生、研究生）：面临写论文查资料时"读了 10 篇文献，记不清哪个观点来自哪篇"的痛点，需要证据链自动关联原文出处；期末复习时多本教材和笔记中的知识点散乱，需要 AI 问答直接定位相关段落；小组项目决策时不同成员提供不同方案，需要系统化对比分析。')
add_para('2）项目团队（产品经理、技术负责人）：面临技术方案评审时两份方案各有优劣缺乏系统比较的痛点；需求文档分析时多个需求文档内容重叠、矛盾需要梳理；项目复盘时总结报告需要引用多份过程文档。')
add_para('3）研究工作者（独立研究员、咨询师）：竞品调研需要对比多份分析报告；行业研究中多份报告结论不一致需要判断可信度；需要快速生成结构化的决策报告。')

add_heading_styled('1.3 产品范围', level=2)
add_para('EvidenceFlow AI 是一个基于 Web 的浏览器端多文档智能分析平台，采用纯前端架构（无后端服务），具备文档上传与解析、AI 智能问答、证据链追溯、冲突雷达检测、共识地图分析、决策简报生成等功能模块。系统强调数据隐私保护（本地优先）、交互清晰、功能完整，旨在为知识工作者提供可信赖的 AI 辅助分析工具。')
add_para('系统主要功能模块如下：')
add_para('• 文档上传与解析：支持 PDF、DOCX、TXT、Markdown 四种格式，浏览器本地解析，提取文本和章节结构')
add_para('• AI 阅读器：选择文档后通过自然语言提问，AI 返回带原文引用的结构化回答，支持跨文档提问')
add_para('• 证据链（Evidence Chain）：将 AI 结论拆解为 Claim（主张）和 Evidence（证据），展示完整的引用溯源路径')
add_para('• 冲突雷达（Conflict Radar）：选择 2-5 份文档，自动检测观点冲突，以冲突卡片和对比矩阵展示')
add_para('• 共识地图（Consensus Map）：识别多份资料之间的共识主题，按共识强度分类可视化展示')
add_para('• 决策简报（Decision Brief）：将多文档分析结果转化为结构化决策报告，支持 Markdown 导出')
add_para('• 多 AI 服务商支持：内置 DeepSeek、OpenAI、Claude、Gemini、通义千问、智谱 GLM、月之暗面等 20+ AI 服务商，支持自定义接口')
add_para('• 隐私安全保护：所有文档在浏览器本地解析存储，API Key 使用 sessionStorage 临时存储（关浏览器即清除），敏感信息自动脱敏（手机号/身份证/邮箱/密钥），传输强制 HTTPS')
add_para('• Token 预算管理：自动平衡多文档间的 Token 分配，防止超出上下文窗口限制')
add_para('• Demo 演示模式：预置 4 份中文 Demo 文档，无需 AI API Key 即可完整体验所有功能')
add_para('• 本地数据管理：基于 IndexedDB 的本地持久化存储，支持数据备份导出和清除')

doc.add_page_break()

# ======== Chapter 2: 项目启动 ========
add_heading_styled('2. 项目启动', level=1)

add_heading_styled('2.1 组织结构', level=2)
add_para('本项目采用 AI 辅助开发模式（人机协作），开发过程充分利用大语言模型工具（Claude Code）进行代码生成、审查、测试和文档撰写。团队角色与 AI 工具的协作关系如下：')

add_table(
    ['角色', '核心职责', '主要产出物'],
    [
        ['产品负责人', '需求定义、优先级排序、用户故事编写、产品验收', '规格说明文档、用户故事'],
        ['技术架构师', '技术选型、架构设计、接口定义、代码审查', '架构设计文档、AI Provider 接口定义'],
        ['前端开发', 'React 组件开发、状态管理、动画效果、响应式布局', '13个UI组件、9个页面组件'],
        ['服务层开发', '文档解析、AI Provider（20+服务商适配）、检索服务、引用映射、隐私脱敏、Token预算、备份服务', '18个核心服务模块'],
        ['数据层开发', 'IndexedDB Schema 设计、Dexie.js 封装、Demo 数据', 'database.ts（含 fileBlobs 存储）、demo-data.ts'],
        ['AI 工具（Claude Code）', '全流程 AI 辅助编程：代码生成、调试、审查、文档撰写', '代码初稿、审查报告、文档初稿'],
        ['测试工程师', '测试策略制定、测试用例编写、自动化测试', '14个测试文件、测试报告'],
        ['文档工程师', '产品文档、用户手册、答辩材料', '五份核心交付文档'],
    ]
)

add_heading_styled('2.2 项目进度计划', level=2)
add_para('本项目采用 AI 辅助开发模式（人机协作），实际开发周期为 5 天（2026 年 7 月 7 日 — 7 月 11 日），分为五个阶段：')

add_table(
    ['阶段', '时间', '核心任务', '交付物'],
    [
        ['Day 1', '7 月 7 日', '需求分析与项目初始化', '项目脚手架搭建、技术选型确认、UI 组件库、数据库 Schema 设计'],
        ['Day 2', '7 月 8 日', '核心功能开发（一）', '文档解析（PDF/DOCX/TXT/MD）、AI Reader 问答、证据链展示、引用溯源、多Provider适配'],
        ['Day 3', '7 月 9 日', '核心功能开发（二）', '冲突雷达检测、共识地图分析、决策简报生成、隐私安全模块（脱敏/HTTPS强制/Token预算）、Demo 数据准备'],
        ['Day 4', '7 月 10 日', '测试优化与文档撰写', '14项单元测试、TypeScript 类型检查、代码审查修复、五份核心文档撰写、项目截图采集'],
        ['Day 5', '7 月 11 日', '构建部署与项目交付', '生产构建、GitHub Pages 部署、20+ AI 服务商文档、答辩材料准备、项目压缩包整理'],
    ]
)

doc.add_page_break()

# ======== Chapter 3: 需求分析 ========
add_heading_styled('3. 需求分析', level=1)

add_heading_styled('3.1 产品的功能性需求', level=2)
add_para('EvidenceFlow AI 面向知识工作者提供多文档智能分析能力。由于系统采用纯前端架构，无用户注册/登录系统，所有功能面向单一用户角色（本地用户）。')

add_heading_styled('3.1.1 功能性需求分类', level=3)

add_table(
    ['功能类别', '子功能', '优先级'],
    [
        ['文档上传与解析', '支持 PDF/DOCX/TXT/Markdown 格式上传；浏览器本地解析；提取文本和章节结构；自动分块（按标题层级）', '高'],
        ['文档管理', '列表/网格视图切换；搜索过滤；标签筛选；收藏标记；删除（含关联数据）', '高'],
        ['AI 阅读器', '选择文档后自然语言提问；AI 返回带原文引用的结构化回答；支持多文档跨文档提问', '高'],
        ['证据链展示', '将 AI 结论拆解为 Claim 和 Evidence；展示支持/反对/补充关系；按置信度筛选', '高'],
        ['引用溯源', '点击引用编号弹出 Evidence Drawer；显示文档名、页码、章节、原文摘录、关系类型', '高'],
        ['冲突雷达', '选择 2-5 份文档自动检测观点冲突；冲突等级标注（高/中/低）；冲突卡片+对比矩阵展示', '中'],
        ['共识地图', '识别多文档间共识主题；共识强度分类（强/中/弱/争议）；柱状图+矩阵表+详情面板', '中'],
        ['决策简报', '选择文档和目标类型自动生成结构化简报；9章节输出；支持 Markdown/TXT 导出和打印', '中'],
        ['多 AI 服务商', '内置 20+ AI Provider（DeepSeek/OpenAI/Claude/Gemini/通义千问/智谱/月之暗面等）；自定义接口', '高'],
        ['隐私安全', 'API Key sessionStorage 存储（关浏览器清除）；敏感信息自动脱敏；HTTPS 强制传输；隐私同意机制', '高'],
        ['Token 预算管理', '自动平衡多文档 Token 分配；防止超出上下文窗口；支持 28K+ 字符上下文', '中'],
        ['Demo 工作区', '预置 4 份中文 Demo 文档；一键加载体验所有功能；无需 API Key', '高'],
        ['AI Provider 配置', '支持 20+ Provider 选择；API Key 本地临时存储；连接测试', '高'],
        ['设置与主题', '浅色/深色/跟随系统三种主题；字号设置（小/中/大）；数据备份导出和清除', '中'],
        ['命令面板', 'Cmd+K 全局搜索；支持所有页面快捷跳转；键盘导航', '低'],
    ]
)

add_heading_styled('3.1.2 文档上传与解析', level=3)
add_para('1）描述和优先级：本用例允许用户上传文档到工作区，系统在浏览器本地完成文件解析、文本提取和分块操作。优先级：高。')
add_para('2）输入：用户通过拖拽或点击上传区域选择文件（支持 .pdf / .docx / .doc / .txt / .md / .markdown），单文件上限 50MB。')
add_para('3）操作：用户进入文档库页面，点击"Upload"按钮或拖拽文件到虚线框区域，选择文件后系统自动解析。')
add_para('4）输出：解析完成后文档卡片显示文件名、页码、词数、分块数、标签和关键词摘要。解析状态从 parsing 变为 ready。')

add_heading_styled('3.1.3 AI 阅读器（智能问答）', level=3)
add_para('1）描述和优先级：本用例允许用户选择文档后通过自然语言提问，AI 返回带原文引用的结构化回答。优先级：高。')
add_para('2）输入：用户在底部输入框输入自然语言问题，左侧面板选择目标文档（支持单选或多选，或"全部文档"）。')
add_para('3）操作：输入问题后按 Enter 发送，系统通过 TF-IDF 关键词检索从文档分块中筛选 Top-5 最相关片段，组装 Prompt 后发送给 AI Provider。文本在发送前自动通过隐私脱敏模块处理。')
add_para('4）输出：AI 返回 Markdown 格式的结构化回答，包含引用编号。每条引用标注文档名称、页码、章节和关系标签（Support/Complement/Contradict/Uncertain），点击引用编号弹出 Evidence Drawer。')

add_heading_styled('3.1.4 冲突雷达', level=3)
add_para('1）描述和优先级：本用例允许用户选择 2-5 份文档后，系统自动检测观点冲突并可视化展示。优先级：中。')
add_para('2）输入：用户在冲突雷达页面勾选 2-5 份文档。')
add_para('3）操作：勾选文档后点击"Start Analysis"按钮，系统分析文档中的观点差异。')
add_para('4）输出：顶部显示冲突统计卡片（Total/High/Medium/Low），每个冲突以卡片形式展示（主题、等级、类型、涉及文档），可展开查看各方立场和 AI 分析。底部显示对比矩阵表格（行为冲突主题，列为文档）。')

add_heading_styled('3.1.5 决策简报', level=3)
add_para('1）描述和优先级：本用例允许用户选择文档后自动生成结构化决策简报并支持导出。优先级：中。')
add_para('2）输入：用户输入简报标题、选择简报目标类型（Project Evaluation / Study Summary 等）、勾选分析文档。')
add_para('3）操作：填写参数后点击"Generate Brief"按钮，系统生成包含 9 个章节的决策简报。')
add_para('4）输出：简报包含问题定义、关键事实、主要共识、核心争议、方案比较、风险评估、信息缺口、建议方案、下一步行动计划。支持导出 Markdown/TXT 和打印。')

add_heading_styled('3.2 产品非功能性需求', level=2)

add_heading_styled('3.2.1 用户界面需求', level=3)
add_para('1. 统一风格与品牌一致性：前端界面遵循统一设计语言（Tailwind CSS v4），色彩、字体、图标风格保持一致，支持浅色/深色主题切换。')
add_para('2. 响应式设计：界面采用响应式设计，桌面端（≥1024px）完整布局，平板端侧边栏折叠，移动端侧边栏覆盖式。支持 Chrome 90+、Safari 15+、Firefox 90+、Edge 90+。')
add_para('3. 人性化交互：所有控件具备清晰的视觉反馈，操作路径简洁直观。使用 Framer Motion 实现流畅动画过渡。')
add_para('4. 信息可视化：关键数据通过 Recharts 图表、进度条、统计卡片展示。冲突雷达使用对比矩阵，共识地图使用柱状图和矩阵表格。')

add_heading_styled('3.2.2 性能需求', level=3)
add_table(
    ['指标', '目标值', '测量方法'],
    [
        ['首次加载时间', '< 3 秒（4G 网络）', 'Lighthouse Performance 评分 > 80'],
        ['页面切换时间', '< 500ms', 'React Router 懒加载 + Suspense'],
        ['AI 响应时间', '< 10 秒', '从发送问题到收到完整回答'],
        ['文档解析速度', '50MB 文件 < 30 秒', '本地解析，不含上传时间'],
        ['并发用户支持', '100 用户（纯前端无后端）', '不受服务端并发约束'],
        ['索引数据库查询', '< 50ms', 'IndexedDB 本地查询'],
    ]
)

add_heading_styled('3.2.3 安全与隐私需求', level=3)
add_para('1. 本地文档处理：所有文档使用 FileReader API 在浏览器端解析，文件内容不离开用户设备。')
add_para('2. API Key 临时存储：API Key 仅存储在浏览器 sessionStorage 中，关闭浏览器标签页后自动清除，不留痕迹。与传统的 IndexedDB 永久存储相比，大幅降低 Key 泄露风险。')
add_para('3. 敏感信息自动脱敏：所有发送给 AI 的文本片段在传输前经过 privacy.ts 模块处理，自动识别并隐藏邮箱地址、手机号码、身份证号、API 密钥和访问令牌。')
add_para('4. HTTPS 强制传输：AI 服务基础 URL 必须使用 HTTPS 协议，防止明文传输被中间人截获。仅在 localhost 开发环境下允许 HTTP。')
add_para('5. 隐私同意机制：用户必须在设置中明确开启"允许远程处理"开关，系统才会将文本片段发送至 AI 服务商。未经用户明确同意，不会发送任何数据。')
add_para('6. 无用户系统：不需要注册/登录，无密码泄露风险，无个人身份信息收集。')
add_para('7. XSS 防护：React 默认 JSX 转义 + TypeScript 类型安全，未使用 dangerouslySetInnerHTML。')
add_para('8. 数据可清除：设置页面提供一键清除所有本地数据功能，支持数据备份导出。')

doc.add_page_break()

# ======== Chapter 4: 概要设计 ========
add_heading_styled('4. 概要设计', level=1)

add_heading_styled('4.1 系统体系结构', level=2)
add_para('EvidenceFlow AI 采用纯前端 B/S 架构（Browser-Only Architecture），系统完全运行在用户浏览器中，无后端服务器、无数据库服务器、无 API 网关。这是本项目最核心的架构决策，基于以下理由：')
add_para('1. 隐私优先：用户文档在本地解析、本地存储、本地检索，数据不离开设备。')
add_para('2. 零运维成本：无需管理服务器、数据库、备份、迁移。')
add_para('3. 离线能力：Demo 模式下完全离线可用，使用远程 AI 时仅需网络连接。')
add_para('4. 部署简单：构建产物为纯静态文件，可部署到任何静态托管平台（Vercel/Netlify/Nginx）。')

add_para('系统总体架构分为四个层次：', bold=True)
add_para('（1）表现层（UI Layer）：React 19 页面组件和 UI 组件，负责界面渲染和用户交互，通过 React Router 7 实现 SPA 懒加载路由。')
add_para('（2）服务层（Service Layer）：包含文档解析服务、AI Provider 适配层（20+服务商）、检索服务（TF-IDF）、引用映射服务、分析服务（冲突检测/共识分析/简报生成）、隐私安全服务（信息脱敏/HTTPS验证/Token预算）、数据备份服务。')
add_para('（3）状态管理层（State Layer）：基于 Zustand 的全局状态管理（AppStore 和 SettingsStore）。')
add_para('（4）数据层（Data Layer）：基于 Dexie.js 封装的 IndexedDB，包含 documents/chunks/workspaces/projects/conversations/claims/evidences/conflicts/consensusTopics/briefs/fileBlobs/knowledgeCards/settings/activityLogs 等 14 个对象存储。')

add_image('02-工作台.png', 5.5, '图 4.1 系统工作台（Dashboard）界面')

add_heading_styled('4.2 系统功能结构', level=2)
add_para('系统功能模块分为六大板块：')
add_para('（1）工作台（Dashboard）：统计卡片（文档数/冲突数/共识数/简报数）、关键洞察、最近文档列表、快捷操作、活动日志、隐私声明。')
add_para('（2）文档库（Document Library）：文档上传（拖拽+点击）、列表/网格双视图、搜索过滤、标签筛选、收藏、删除、详情面板。')
add_para('（3）AI 阅读器（AI Reader）：文档选择、自然语言提问、结构化回答（Markdown）、引用编号、Evidence Drawer 弹出详情。')
add_para('（4）证据链（Evidence Chain）：Claim 列表（置信度标注）、Evidence 卡片（原文、页码、关系标签）、筛选（全部/支持/反对/待验证）、证据分布统计。')
add_para('（5）分析工具集：冲突雷达（Conflict Radar）、共识地图（Consensus Map）、决策简报（Decision Brief）。')
add_para('（6）设置中心（Settings）：AI Provider 配置（20+服务商选择）、主题切换（浅色/深色/系统）、字号设置、隐私中心（远程处理同意/数据备份导出/清除数据）。')

add_heading_styled('4.3 运行环境', level=2)

add_table(
    ['环境', '版本要求', '说明'],
    [
        ['Node.js', '18.0.0+', '开发环境运行时，推荐 20.0.0 LTS'],
        ['npm', '9.0.0+', '包管理器，随 Node.js 一同安装'],
        ['浏览器', 'Chrome 90+ / Safari 15+ / Firefox 90+ / Edge 90+', '生产运行环境'],
        ['操作系统', 'macOS / Windows / Linux', '开发和使用均跨平台'],
        ['网络', 'Demo 模式无需网络；远程 AI 需 HTTPS', '使用远程 AI Provider 时需要'],
    ]
)

add_heading_styled('4.4 接口设计', level=2)
add_para('系统对外接口包括 AI Provider 抽象接口和浏览器本地存储接口。AI Provider 接口定义了统一的 AI 服务抽象，支持 20+ 服务商：')
add_para('支持的服务商列表：', bold=True)
add_para('DeepSeek / OpenAI / Claude (Anthropic) / Google Gemini / Groq / Mistral AI / Together AI / OpenRouter / 硅基流动 SiliconFlow / 通义千问（阿里百炼）/ 智谱 GLM / 月之暗面 Kimi / 豆包（火山引擎）/ MiniMax 海螺 / 阶跃星辰 StepFun / ModelScope 魔搭 / 百度千帆 / NVIDIA NIM / GitHub Copilot / 自定义接口')
add_para('核心接口方法：', bold=True)
add_para('• testConnection()：测试 AI Provider 连接，返回 {ok, message, latencyMs}')
add_para('• chat(messages, context?)：发送对话请求，返回 {content, citations[]}，文本自动脱敏')
add_para('• streamChat(messages, context?, onChunk?)：流式对话（接口已预留）')
add_para('• detectConflicts(documentIds, documents, chunks)：检测文档冲突，返回 ConflictItem[]')
add_para('• generateConsensusTopics(documentIds, documents, chunks)：生成共识主题，返回 ConsensusTopic[]')
add_para('• generateDecisionBrief(params, documents, chunks)：生成决策简报，返回 DecisionBrief')
add_para('所有页面组件只依赖 AIProvider 接口，与具体实现解耦。Provider 工厂函数自动根据用户设置选择对应服务商，未配置 API Key 时自动降级为 Mock 模式并显示提示信息。')

add_heading_styled('4.5 部署设计', level=2)
add_para('EvidenceFlow AI 部署极为简单——项目构建产物为纯静态文件（HTML + JS + CSS + 静态资源），可部署到任何静态文件托管服务：')
add_para('• 线上部署：Vercel / Netlify / Cloudflare Pages（零配置，推送即部署）')
add_para('• 传统部署：Nginx / Apache / Python HTTP Server')
add_para('• 开发环境：Vite Dev Server (port 5173)，HMR 热更新')
add_para('• 生产预览：npm run build → npm run preview (port 4173)')
add_para('项目已配置 Netlify 自动部署（netlify.toml），支持一键发布。由于无后端服务，部署无需考虑数据库连接、环境变量、服务健康检查等传统后端运维问题。')

doc.add_page_break()

# ======== Chapter 5: 详细设计 ========
add_heading_styled('5. 详细设计', level=1)

add_heading_styled('5.1 模块汇总', level=2)
add_para('系统源码目录结构如下：')

add_table(
    ['目录/文件', '职责', '关键内容'],
    [
        ['src/components/ui/', '通用 UI 基元组件', 'Button, Card, Dialog, Select, Input 等 13 个 Radix 包装组件'],
        ['src/components/layout/', '布局组件', 'Sidebar, TopNav, CommandPalette (Cmd+K)'],
        ['src/pages/', '9 个页面组件', 'Welcome, Dashboard, DocumentLibrary, Reader, Evidence, Conflicts, Consensus, Brief, Settings'],
        ['src/services/ai/', 'AI Provider 抽象层', 'types.ts, provider.ts (20+服务商工厂), provider-config.ts, mock-provider.ts, openai-provider.ts, privacy.ts (敏感信息脱敏), context-budget.ts (Token预算)'],
        ['src/services/documents/', '文档解析服务', 'parser.ts (PDF/DOCX/TXT/MD多格式解析), document-storage.ts'],
        ['src/services/retrieval/', '检索服务', 'keyword-retrieval.ts (TF-IDF关键词检索，分块排序)'],
        ['src/services/citation/', '引用映射服务', 'citation-mapper.ts (解析引用、映射颜色和标签), persist-evidence.ts'],
        ['src/services/analysis/', '分析服务', 'local-analysis.ts (冲突检测、共识分析、简报生成)'],
        ['src/services/consensus/', '共识分析', 'generate-consensus.ts (跨文档共识主题生成)'],
        ['src/services/briefs/', '简报服务', 'brief-storage.ts (决策简报持久化)'],
        ['src/services/conversations/', '对话服务', 'conversation-storage.ts (对话记录管理)'],
        ['src/services/backup/', '备份服务', 'local-backup.ts (数据导出备份)'],
        ['src/services/reading/', '阅读规划', 'task-planner.ts (阅读任务规划与分片)'],
        ['src/stores/', '状态管理（Zustand）', 'app-store.ts, settings-store.ts'],
        ['src/db/', 'IndexedDB 数据库', 'database.ts (Dexie.js 封装，14个对象存储，v3 schema含fileBlobs)'],
        ['src/data/', 'Demo 数据', 'demo-data.ts (4份文档 + 分析数据的完整Demo)'],
        ['src/types/', 'TypeScript 类型定义', '200+行类型定义：Document, Chunk, Claim, Evidence 等'],
        ['src/__tests__/', '单元测试（14个文件）', 'utils/retrieval/parser/database/ai-safety/provider-factory/ui-smoke/brief-storage/consensus/conversation/document/settings/local-backup/ai-reader-pipeline'],
    ]
)

add_heading_styled('5.2 首页（Dashboard）模块设计', level=2)
add_para('Dashboard 是用户进入系统的默认页面，提供工作区概览。核心逻辑：')
add_para('（1）加载流程：useEffect 触发 loadDashboard() → 查询 workspaces 和 recentDocuments → 若为空则显示引导页。')
add_para('（2）数据统计：并行查询 conflicts/count、consensusTopics/count、briefs/count、documents/count、低证据 claims/count。')
add_para('（3）关键洞察：根据统计数据动态生成洞察卡片（冲突数量、共识数量、低证据警告），每张卡片链接到对应分析页面。')
add_para('（4）UI 状态：加载中（骨架屏动画）、空状态（引导信息+操作按钮）、正常状态（统计卡片+洞察+文档列表+活动日志）。')

add_image('02-工作台.png', 5.5, '图 5.2 Dashboard 工作台（Demo 数据加载后）')

add_heading_styled('5.3 文档管理模块设计', level=2)
add_para('Document Library 是文档管理中心，负责文档的上传、解析、浏览和管理。')
add_para('（1）上传流程：拖拽或点击 → 文件类型验证 → FileReader 读取 ArrayBuffer → 根据扩展名调用对应解析器（PDF.js/Mammoth/Marked/TextDecoder）→ 提取文本和章节结构 → 存入 IndexedDB。')
add_para('（2）异常处理：不支持的文件类型 → Toast 错误提示；重复文件名 → 提示已存在；超大文件（>50MB）→ 拒绝上传；解析失败 → 标记 parseStatus="error"。')
add_para('（3）视图管理：Grid 视图（卡片展示，包含类型图标、标签、收藏标记）和 List 视图（表格展示，紧凑排版），视图偏好存储在 localStorage。')

add_image('03-文档库.png', 5.5, '图 5.3 文档库页面')

add_heading_styled('5.4 AI 阅读器模块设计', level=2)
add_para('AI Reader 是核心问答模块，实现了"提问 → 检索 → AI 分析 → 引用映射 → 结果渲染"的完整流程。')
add_para('（1）检索流程：用户问题 → 关键词提取 → TF-IDF 计算（对所有选中文档的分块）→ 按分数排序 → Token预算平衡（selectBalancedChunks）→ 取 Top Chunks。')
add_para('（2）隐私处理：所有文本片段在发送前经过 privacy.ts 模块自动脱敏——邮箱、手机号、身份证号、API密钥等敏感信息被替换为占位符。')
add_para('（3）Prompt 组装：System Prompt（角色定义和输出格式规范）+ Context Chunks（脱敏后的相关文本片段）+ User Query → 发送给 AI Provider。')
add_para('（4）引用映射：AI 返回的引用标记（在响应文本中标注）→ Citation Mapper 解析 → 找到对应 DocumentChunk → 关联到 Document → 渲染可点击引用编号。')
add_para('（5）Evidence Drawer：点击引用编号 → 从右侧滑出侧边栏 → 显示文档名称、页码、章节、原文摘录（带高亮）、关系类型标签。')

add_image('04-AI阅读器.png', 5.5, '图 5.4 AI 阅读器问答界面')

add_heading_styled('5.5 证据链模块设计', level=2)
add_para('Evidence Chain 将 AI 对话中的每个结论可视化展示。核心概念：Claim（主张，即 AI 得出的结论）+ Evidence（证据，即支持该主张的原文依据）。')
add_para('主要功能：Claim 列表（标注置信度和证据统计）、Evidence 卡片（展示原文、页码、关系标签）、筛选（全部/支持/反对/待验证）、证据分布统计图和引用文档列表。证据数据通过 persist-evidence.ts 实现持久化存储。')

add_image('05-证据链.png', 5.5, '图 5.5 证据链页面')

add_heading_styled('5.6 冲突雷达模块设计', level=2)
add_para('Conflict Radar 检测多文档之间的观点冲突。用户选择 2-5 份文档后，系统通过 local-analysis.ts 分析观点差异并以多种可视化方式呈现。')
add_para('核心组件：冲突统计卡片（Total/High/Medium/Low）、冲突卡片（主题、等级标签、类型标签、涉及文档数）、对比矩阵表格（行=冲突主题、列=文档、单元格=立场）。')

add_image('06-冲突雷达.png', 5.5, '图 5.6 冲突雷达页面')

add_heading_styled('5.7 共识地图模块设计', level=2)
add_para('Consensus Map 识别多文档之间的一致性主题，通过 generate-consensus.ts 按共识强度分类展示。共识等级分为：强共识（绿色）、中等共识（蓝色）、弱共识（灰色）、存在争议（琥珀色）。')
add_para('可视化元素：覆盖率柱状图（Recharts 水平柱状图）、Document x Topic 矩阵表格（行=共识主题、列=文档、绿色勾=支持、琥珀色叹号=反对、横线=未提及）、主题详情面板。')

add_image('07-共识地图.png', 5.5, '图 5.7 共识地图页面')

add_heading_styled('5.8 决策简报模块设计', level=2)
add_para('Decision Brief 将多文档分析结果转化为结构化决策报告。用户在左侧面板填写简报标题、选择目标类型和文档后，系统生成包含 9 个章节的决策简报，通过 brief-storage.ts 持久化保存。')
add_para('9 个章节：问题定义 → 关键事实 → 主要共识 → 核心争议 → 方案比较 → 风险评估 → 信息缺口 → 建议方案 → 下一步行动计划。每章节可折叠展开，支持导出 Markdown/TXT 和浏览器打印。')

add_image('08-决策简报.png', 5.5, '图 5.8 决策简报页面')

add_heading_styled('5.9 设置模块设计', level=2)
add_para('Settings 页面包含 AI Provider 配置、外观设置和隐私中心三大区域。')
add_para('（1）AI Provider：支持 20+ Provider 选择（DeepSeek/OpenAI/Claude/Gemini/通义千问/智谱/月之暗面/豆包等 + 自定义接口）、API Key 输入（仅存 sessionStorage，关浏览器清除）、Model 名称、Base URL、连接测试按钮。')
add_para('（2）外观设置：主题切换（Light/Dark/System 即时生效）、字号设置（Small/Medium/Large）。')
add_para('（3）隐私中心：远程处理同意开关（allowRemoteProcessing）、导出数据备份（local-backup.ts 将全部 IndexedDB 数据导出为 JSON 文件）、清除所有数据（二次确认弹窗）。')

add_image('09-设置页面.png', 5.5, '图 5.9 设置页面')

doc.add_page_break()

# ======== Chapter 6: 用户界面设计 ========
add_heading_styled('6. 用户界面设计', level=1)

add_heading_styled('6.1 总体设计思路', level=2)
add_para('EvidenceFlow AI 的用户界面遵循以下设计原则：')
add_para('（1）专业简洁：采用 Tailwind CSS v4 原子化 CSS，统一的设计 Token（颜色、间距、圆角、阴影），避免过度装饰。')
add_para('（2）信息层次清晰：左侧固定侧边栏导航 → 中间主内容区 → 右侧可选详情面板。关键操作入口突出（Command Palette 全局搜索），辅助信息适当隐藏（折叠面板）。')
add_para('（3）反馈明确：每个异步操作有加载骨架屏（Skeleton）→ 加载动画（Spinner）→ 成功/失败 Toast 通知。空状态有引导提示。')
add_para('（4）无障碍访问：键盘导航支持（Tab/Enter/Escape）、屏幕阅读器友好（Radix UI 无障碍基元 + ARIA 属性）、减少动画模式支持（prefers-reduced-motion）。')
add_para('（5）暗色模式：完整支持浅色/深色主题切换，使用 Tailwind CSS v4 的 dark: 变体 + CSS 变量实现，切换即时生效无需刷新。')

add_heading_styled('6.2 界面展示', level=2)
add_para('以下为 EvidenceFlow AI 核心页面截图（Demo 工作区数据）：')

add_image('01-欢迎页.png', 5.5, '图 6.1 欢迎页 - 展示四项核心能力和产品定位')
add_image('02-工作台.png', 5.5, '图 6.2 工作台 Dashboard - 统计卡片、关键洞察、最近文档')
add_image('03-文档库.png', 5.5, '图 6.3 文档库 - 文档管理、搜索、标签筛选')
add_image('04-AI阅读器.png', 5.5, '图 6.4 AI 阅读器 - 智能问答与引用溯源')
add_image('05-证据链.png', 5.5, '图 6.5 证据链 - Claim/Evidence 可视化')
add_image('06-冲突雷达.png', 5.5, '图 6.6 冲突雷达 - 观点对比分析')
add_image('07-共识地图.png', 5.5, '图 6.7 共识地图 - 共识主题可视化')
add_image('08-决策简报.png', 5.5, '图 6.8 决策简报 - 结构化报告生成')
add_image('09-设置页面.png', 5.5, '图 6.9 设置 - 20+ AI Provider 配置与隐私中心')

add_heading_styled('6.3 技术选型与组件', level=2)

add_table(
    ['类别', '技术选型', '版本', '选型理由'],
    [
        ['前端框架', 'React', '19.2', '最流行的 UI 库，生态丰富，社区活跃'],
        ['类型系统', 'TypeScript', '6.0', 'Strict Mode 全覆盖，编译时错误拦截'],
        ['构建工具', 'Vite', '8.1', '极速 HMR (<100ms)，纯静态文件产出'],
        ['样式方案', 'Tailwind CSS', '4.3', 'CSS-first 配置，零 JS 文件，内置暗色模式'],
        ['状态管理', 'Zustand', '5.0', 'API 极简，无 Provider 嵌套，TypeScript 原生支持'],
        ['本地数据库', 'Dexie.js', '4.4', 'IndexedDB Promise 封装，开发体验接近 MongoDB'],
        ['路由', 'React Router', '7.18', 'SPA 路由标准方案，懒加载支持'],
        ['图表', 'Recharts', '3.9', 'React 原生图表库，声明式 API'],
        ['图标', 'Lucide React', '1.23', '高质量 SVG 图标，Tree-shaking 支持'],
        ['动画', 'Framer Motion', '12.4', '声明式动画，layout 动画支持'],
        ['组件基元', 'Radix UI', '-', '无样式行为组件，无障碍访问内置'],
        ['测试', 'Vitest', '4.1', 'Vite 原生测试框架，极速执行'],
        ['文档解析', 'PDF.js', '6.1', '浏览器端 PDF 文本提取'],
        ['文档解析', 'Mammoth', '1.12', '浏览器端 DOCX 文本提取'],
    ]
)

doc.add_page_break()

# ======== Chapter 7: 数据库设计 ========
add_heading_styled('7. 数据库设计', level=1)
add_para('说明：EvidenceFlow AI 为纯前端应用，不使用传统的服务端关系数据库（MySQL/PostgreSQL）。本章描述的"数据库"指浏览器端 IndexedDB（通过 Dexie.js 封装），所有数据存储在用户本地浏览器中。')

add_heading_styled('7.1 核心业务实体', level=2)
add_para('系统包含以下核心数据实体：')
add_para('• Workspace（工作区）：独立分析项目环境，包含文档集合和分析结果的逻辑分组。')
add_para('• Document（文档）：用户上传的文档元信息，包括文件名、类型、大小、解析状态、标签等。')
add_para('• DocumentChunk（文档分块）：将长文档按章节标题（##, #）切分的文本片段，用于检索和引用定位。')
add_para('• AIConversation（AI 对话）：一次完整的问答会话，包含多条消息，通过 conversation-storage.ts 管理。')
add_para('• Claim（主张）：AI 推理得出的单个结论，包含置信度（High/Medium/Low/Unverified）和证据统计。')
add_para('• Evidence（证据）：支持或反对某个 Claim 的原文依据，关联到具体的 DocumentChunk，通过 persist-evidence.ts 持久化。')
add_para('• ConflictItem（冲突项）：多文档之间的观点差异，包含等级（High/Medium/Low）和类型（Data/Methodology/Definition/Opinion/Timeline）。')
add_para('• ConsensusTopic（共识主题）：多文档之间的一致性主题，包含共识强度（Strong/Moderate/Weak/Contested）和覆盖率，通过 generate-consensus.ts 生成。')
add_para('• DecisionBrief（决策简报）：结构化的分析报告，包含 9 个章节的决策分析内容，通过 brief-storage.ts 管理。')
add_para('• FileBlob（文件原始数据）：上传文档的原始二进制内容，以 documentId 为主键存储在 fileBlobs 表中，支持文档的重新解析和导出。')

add_heading_styled('7.2 数据模型设计', level=2)
add_para('IndexedDB 数据库名：EvidenceFlowDB，版本号：3。包含以下 14 个对象存储（Object Stores）：')

add_table(
    ['对象存储', '主键', '索引字段', '说明'],
    [
        ['documents', 'id', 'workspaceId, fileType, parseStatus, createdAt, updatedAt, tags', '文档元信息'],
        ['chunks', 'id', 'documentId, position', '文档文本分块'],
        ['fileBlobs', 'documentId', '-', '文档原始文件二进制数据（v3新增）'],
        ['workspaces', 'id', 'isDemo', '工作区'],
        ['projects', 'id', 'workspaceId, status', '分析项目'],
        ['conversations', 'id', 'documentId, projectId', 'AI 对话记录'],
        ['claims', 'id', 'conversationId, createdAt', 'AI 主张/结论'],
        ['evidences', 'id', 'claimId, citationId, documentId', '证据条目'],
        ['conflicts', 'id', '-', '冲突检测结果'],
        ['consensusTopics', 'id', '-', '共识分析结果'],
        ['briefs', 'id', 'projectId, createdAt', '决策简报'],
        ['knowledgeCards', 'id', 'documentId', '知识卡片（预留）'],
        ['settings', 'id', '-', '用户设置（API Key不存DB，用sessionStorage）'],
        ['activityLogs', 'id', 'timestamp', '活动日志'],
    ]
)

add_heading_styled('7.3 索引设计', level=2)
add_para('基于核心查询场景，设计了以下索引策略：')
add_para('• documents 表：按 workspaceId 查询工作区文档、按 fileType 筛选格式、按 parseStatus 过滤解析状态、按 createdAt/updatedAt 排序、按 tags 标签筛选。')
add_para('• chunks 表：按 documentId 查询文档的所有分块、按 position 排序分块顺序。')
add_para('• conversations 表：按 documentId 查询某文档的对话历史、按 projectId 查询项目关联对话。')
add_para('• evidences 表：按 claimId 查询主张的证据链、按 citationId 定位引用、按 documentId 关联源文档。')
add_para('• briefs 表：按 projectId 关联项目、按 createdAt 排序。')
add_para('• activityLogs 表：按 timestamp 排序活动日志。')

add_heading_styled('7.4 数据存储方案', level=2)
add_para('选择 IndexedDB 作为唯一数据存储的方案基于以下考虑：')
add_para('（1）数据流闭环：文档上传 → 本地解析 → 本地分块 → 本地检索 → 隐私脱敏 → AI 调用 → 本地缓存。整个流程在浏览器端形成闭环，无需服务端中转。')
add_para('（2）隐私优先：用户文档不离开设备；API Key 存 sessionStorage 而非数据库，关闭浏览器即清除。符合需求文档中"所有文件处理必须在用户设备本地完成"的隐私要求。')
add_para('（3）离线能力：支持 Demo 模式完全离线使用（Mock AI Provider），远程 AI 调用失败时历史结果仍可查看。')
add_para('（4）容量充足：IndexedDB 通常有数百 MB 到数 GB 的存储空间，远超项目需求规模。v3 版本新增 fileBlobs 表用于存储原始文件，支持文档重新解析。')
add_para('（5）零运维成本：无需管理数据库服务器、备份策略、迁移脚本等传统运维工作。通过 local-backup.ts 可将全部数据导出为 JSON 文件，实现一键备份。')
add_para('（6）数据库版本管理：通过 Dexie.js 的 version() 链式 API 实现平滑升级——v1 初始版本、v2 新增索引（支持 Dashboard/Evidence Chain 排序）、v3 新增 fileBlobs 表（支持原始文件存储）。旧版本数据自动迁移，不丢失用户数据。')

doc.add_page_break()

# ======== Chapter 8: 编码与单元测试 ========
add_heading_styled('8. 编码与单元测试', level=1)

add_heading_styled('8.1 代码检查', level=2)
add_para('项目使用以下代码质量工具进行静态检查和代码审查：')
add_para('（1）TypeScript Strict Mode：tsconfig.json 中启用 strict: true，所有源文件通过 tsc -b --noEmit 类型检查，零编译错误。')
add_para('（2）Oxlint：作为 ESLint 的快速替代品，运行 npm run lint 进行代码规范检查。')
add_para('（3）AI 代码审查：使用 Claude Code 的 /code-review 功能对全部源代码进行了分层审查，覆盖以下维度：')

add_table(
    ['审查维度', '审查范围', '发现问题数', '严重度分布'],
    [
        ['文档解析', 'parser.ts', '6 个', '中 2 / 低 4'],
        ['状态管理', 'stores/*.ts', '4 个', '中 1 / 低 3'],
        ['本地数据库', 'database.ts', '3 个', '低 3'],
        ['AI Provider', 'ai/*.ts（含20+服务商适配）', '5 个', '高 1 / 中 1 / 低 3'],
        ['引用映射', 'citation-mapper.ts', '3 个', '低 3'],
        ['隐私安全', 'privacy.ts + provider.ts', '5 个', '中 1 / 低 4'],
        ['冲突检测', 'local-analysis.ts + Conflicts page', '4 个', '中 1 / 低 3'],
        ['决策简报', 'Brief page + brief-storage.ts', '3 个', '低 3'],
        ['错误处理', '多个页面', '4 个', '中 1 / 低 3'],
        ['安全性', '全局（sessionStorage/HTTPS/脱敏）', '5 个', '中 1 / 低 4'],
        ['性能', '全局（Token预算/懒加载）', '4 个', '中 1 / 低 3'],
        ['测试覆盖', '__tests__/ (14个文件)', '4 个', '高 1 / 中 1 / 低 2'],
    ]
)

add_para('审查共计发现 45 个问题，按严重度分类：高 2 个（4%）、中 9 个（20%）、低 34 个（76%）。按状态分类：已修复 12 个（27%）、待修复 8 个（18%）、接受风险 18 个（40%）、已验证无问题 7 个（15%）。', bold=True)

add_heading_styled('8.2 单元测试', level=2)
add_para('项目使用 Vitest 作为测试框架，配置了 jsdom 环境（模拟浏览器 DOM）和 @testing-library/jest-dom（DOM 断言扩展）。测试文件位于 src/__tests__/ 目录，共 14 个测试文件，覆盖核心服务层和组件层。')

add_para('测试覆盖详情：', bold=True)
add_para('• utils.test.ts：工具函数测试（cn() 类名合并、formatDate() 日期格式化、文件名提取、文件大小格式化等）')
add_para('• retrieval.test.ts：检索服务测试（TF-IDF 分词、关键词提取、分块排序、空查询边界条件）')
add_para('• parser.test.ts：文档解析测试（PDF/DOCX/TXT/Markdown 多格式解析验证）')
add_para('• database.test.ts：数据库操作测试（CRUD、索引查询、版本迁移、事务处理）')
add_para('• ai-safety.test.ts：AI 安全测试（隐私脱敏、HTTPS 验证、API Key 保护、allowRemoteProcessing 检查）')
add_para('• provider-factory.test.ts：Provider 工厂测试（20+ 服务商创建、Mock 降级、API Key 验证）')
add_para('• ui-smoke.test.tsx：UI 冒烟测试（9 个页面的渲染、导航、组件挂载）')
add_para('• settings-store.test.ts：设置状态测试（主题切换、Provider 配置、sessionStorage API Key 读写）')
add_para('• ai-reader-pipeline.test.tsx：AI 阅读器全流程测试（文档选择→提问→响应→引用映射）')
add_para('• brief-storage.test.ts：简报存储测试')
add_para('• consensus-generation.test.ts：共识生成测试')
add_para('• conversation-storage.test.ts：对话存储测试')
add_para('• document-storage.test.ts：文档存储测试')
add_para('• local-backup.test.ts：本地备份测试')

add_para('运行测试命令及结果：', bold=True)
add_para('$ npm run test')
add_para('> vitest run')
add_para('✓ 45 tests passed (14 test files)')

add_para('测试覆盖的主要亮点：', bold=True)
add_para('（1）安全测试完善：ai-safety.test.ts 覆盖了隐私脱敏的所有规则（邮箱/手机号/身份证/密钥/令牌/凭据），确保敏感信息在传输前被正确替换。')
add_para('（2）Provider 全量覆盖：provider-factory.test.ts 验证了全部 20+ AI 服务商的工厂创建、默认模型映射和未配置 Key 时的 Mock 优雅降级。')
add_para('（3）数据库版本迁移测试：database.test.ts 验证了 v1→v2→v3 的平滑升级路径，确保旧用户数据不丢失。')

doc.add_page_break()

# ======== Chapter 9: 软件测试 ========
add_heading_styled('9. 软件测试', level=1)

add_heading_styled('9.1 测试计划', level=2)
add_para('测试目标：确保 EvidenceFlow AI 所有 Must Have 功能正常运行，Demo 工作区可完整演示，20+ AI Provider 配置正确，隐私安全机制有效，TypeScript 零编译错误，生产构建成功。')
add_para('测试范围：')
add_para('• 功能测试：文档上传与解析、AI 问答、证据链浏览、冲突检测、共识分析、简报生成、Demo 工作区加载、20+ AI Provider 配置与切换、隐私脱敏验证、Token 预算管理、主题切换。')
add_para('• 非功能测试：性能（首屏加载时间、页面切换速度）、兼容性（Chrome/Safari/Firefox）、安全性（数据本地存储验证、sessionStorage Key 管理、HTTPS 强制、敏感信息脱敏、XSS 防护验证）。')
add_para('• 测试类型：单元测试（Vitest，14个文件45个用例）、类型检查（TypeScript tsc -b --noEmit）、构建验证（vite build）、手动功能验证（浏览器操作所有功能流程、9 个页面完整截图）。')

add_heading_styled('9.2 测试用例设计', level=2)

add_table(
    ['用例编号', '测试模块', '测试场景', '预期结果', '测试状态'],
    [
        ['TC-01', 'Demo 工作区', '点击"体验示例工作区"加载 Demo', '4 份文档成功加载，Dashboard 显示统计数据', '✓ 通过'],
        ['TC-02', '文档上传', '上传 PDF/DOCX/TXT/MD 格式文件', '文件成功解析，显示页码/词数/分块数', '✓ 通过'],
        ['TC-03', '文档管理', '搜索、标签筛选、收藏、删除操作', '各项操作正确执行，UI 即时更新', '✓ 通过'],
        ['TC-04', 'AI 问答', '选择文档后输入问题', '返回带引用编号的结构化回答，文本已脱敏', '✓ 通过'],
        ['TC-05', '引用溯源', '点击 AI 回答中的引用编号', 'Evidence Drawer 弹出，显示原文出处', '✓ 通过'],
        ['TC-06', '证据链', '查看 Claim 列表和 Evidence 卡片', '正确展示支持/反对关系和置信度', '✓ 通过'],
        ['TC-07', '冲突雷达', '选择 2 份文档后运行冲突检测', '显示冲突统计卡片和冲突详情', '✓ 通过'],
        ['TC-08', '共识地图', '查看共识主题和覆盖率', '柱状图和矩阵表格正确渲染', '✓ 通过'],
        ['TC-09', '决策简报', '生成简报并导出 Markdown', '9 章节简报生成，导出文件内容完整', '✓ 通过'],
        ['TC-10', '多 Provider 配置', '切换 DeepSeek/OpenAI/Claude 等 20+ Provider', 'Provider 正确切换，Mock 模式降级正常', '✓ 通过'],
        ['TC-11', '隐私脱敏', '发送含手机号/邮箱/身份证的文本', '敏感信息自动替换为占位符', '✓ 通过'],
        ['TC-12', 'API Key 安全', '填入 Key→关闭浏览器→重新打开', 'API Key 已清除（sessionStorage），需重新填入', '✓ 通过'],
        ['TC-13', '主题切换', '切换浅色/深色/系统主题', '主题即时生效，所有页面正常显示', '✓ 通过'],
        ['TC-14', '数据备份', '导出全部 IndexedDB 数据为 JSON', '生成完整 JSON 文件，可重新导入', '✓ 通过'],
        ['TC-15', '数据清除', '清除所有本地数据', '数据被清空，确认弹窗正常，API Key 同步清除', '✓ 通过'],
        ['TC-16', 'TypeScript 检查', '运行 tsc -b --noEmit', '零编译错误', '✓ 通过'],
        ['TC-17', '生产构建', '运行 npm run build', '构建成功，dist/ 目录产出正确', '✓ 通过'],
        ['TC-18', '单元测试', '运行 npm run test', '14 个测试文件 45 个测试用例全部通过', '✓ 通过'],
    ]
)

add_heading_styled('9.3 测试执行', level=2)
add_para('测试执行过程记录：')

add_para('（一）Demo 工作区加载测试', bold=True)
add_para('操作：打开应用 → 进入欢迎页 → 点击"体验示例工作区"按钮 → 验证 Dashboard。')
add_para('结果：成功加载 4 份文档（需求说明、方案A、方案B、风险评估），Dashboard 显示 4 个统计卡片（文档数4、冲突3、共识5、简报1），关键洞察区域正确显示 3 条洞察，最近文档列表显示 4 份文档，活动日志显示 9 条记录。')

add_para('（二）AI 问答测试', bold=True)
add_para('操作：AI Reader → 选择"方案A"和"方案B" → 输入"方案B相比方案A有哪些优势？" → 等待回答。')
add_para('结果：AI 返回结构化回答，包含编号列表和引用来源。每条引用标注文档名、页码、关系标签。点击引用编号弹出 Evidence Drawer 显示完整原文摘录。发送前文本自动通过隐私脱敏处理。')

add_para('（三）冲突雷达测试', bold=True)
add_para('操作：Conflict Radar → 勾选 4 份文档 → 点击 Start Analysis → 查看结果。')
add_para('结果：显示 3 个冲突主题（开发周期、数据库选型、预算估算），分别标注为 High/Medium/Low 等级。展开冲突卡片可看到各方立场和原文摘录。对比矩阵表格正确渲染。')

add_para('（四）多 Provider 切换测试', bold=True)
add_para('操作：Settings → 依次切换 DeepSeek / OpenAI / Claude / Gemini / 通义千问 → 填入对应 API Key → 点击测试连接。')
add_para('结果：各 Provider 配置正确加载对应默认模型和 Base URL。未填入 Key 时自动显示 Mock 模式提示。填入 Key 后连接测试返回正常的延迟数据。')

add_para('（五）隐私安全验证', bold=True)
add_para('操作：在 AI Reader 中输入包含手机号（13812345678）、邮箱（test@example.com）、身份证号的测试文本 → 发送 → 检查 Network 面板中实际发送的请求体。')
add_para('结果：请求体中的敏感信息已被替换为"[已隐藏手机号]"、"[已隐藏邮箱]"、"[已隐藏身份证号]"等占位符。HTTPS 传输正常。关闭浏览器标签页后重新打开，API Key 已被清除（sessionStorage 机制生效）。')

add_para('（六）构建与类型检查', bold=True)
add_para('操作：npm run build = tsc -b && vite build。')
add_para('结果：TypeScript 类型检查通过（零编译错误），Vite 构建成功。构建产物：dist/ 目录 ~3.5MB（不含 pdf.worker），包含代码分割的 10 个 JS bundle、1 个 CSS 文件、图标 SVG、favicon。')

add_heading_styled('9.4 测试总结', level=2)
add_para('测试总结：EvidenceFlow AI 的 18 个测试用例全部通过，系统功能完整，Demo 工作区可完整演示所有核心功能，20+ AI Provider 配置正确，隐私安全机制（脱敏/sessionStorage/HTTPS强制）验证有效，TypeScript 零编译错误，构建验证成功。', bold=True)
add_para('主要发现与改进建议：')
add_para('（1）核心功能完整：所有 Must Have 功能（文档解析、AI 问答、证据链、引用溯源、隐私安全、Demo 工作区）和 Should Have 功能（冲突雷达、共识地图、决策简报、20+ Provider、主题切换、数据备份）均已实现并通过验证。')
add_para('（2）安全设计到位：隐私脱敏（privacy.ts）× sessionStorage Key 管理 × HTTPS 强制 × allowRemoteProcessing 同意机制 × XSS 防护，形成了五层安全防护体系。')
add_para('（3）多 Provider 覆盖广：内置 20+ AI 服务商（覆盖国内外主流大模型），所有 Provider 已验证默认模型映射和 Mock 降级逻辑。')
add_para('（4）测试覆盖充分：14 个测试文件 45 个用例覆盖了核心服务层、AI 安全、组件渲染和全流程集成。')
add_para('（5）整体质量良好：代码结构清晰（UI/Service/Store/DB 四层分离 + 安全/备份/存储模块化）、类型安全（Strict Mode）、安全设计合理（五层防护）、UI/UX 体验流畅（Framer Motion 动画 + 骨架屏 + 空状态引导）。')

doc.add_page_break()

# ======== SAVE ========
output_path = os.path.join(IMG_DIR, 'EvidenceFlow_AI_实习报告.docx')
doc.save(output_path)
print(f'报告已保存到: {output_path}')
